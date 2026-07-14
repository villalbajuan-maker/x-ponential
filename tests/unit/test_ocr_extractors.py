from __future__ import annotations

from pathlib import Path

import pandas as pd
import pytest
from docx import Document
from PIL import Image
from openpyxl import Workbook

from business_bridge.adapters import ocr_extractors as oe
from business_bridge.adapters.ocr_extractors import (
    TesseractNotFoundError,
    extract_text_from_csv,
    extract_text_from_docx,
    extract_text_from_file,
    extract_text_from_image,
    extract_text_from_pdf,
    extract_text_from_xlsx,
    preprocess_image_for_ocr,
    run_ocr,
)


class FakePdfDocument:
    def __init__(self, pages: list[object]) -> None:
        self._pages = pages

    def __enter__(self) -> "FakePdfDocument":
        return self

    def __exit__(self, exc_type, exc, tb) -> None:
        return None

    def __iter__(self):
        return iter(self._pages)


def test_preprocess_image_and_run_ocr_success(monkeypatch: pytest.MonkeyPatch) -> None:
    image = Image.new("RGB", (2, 1), color="white")
    image.putpixel((0, 0), (0, 0, 0))

    monkeypatch.setattr(
        "business_bridge.adapters.ocr_extractors.pytesseract.image_to_string",
        lambda img, config=None: "  hola  \r\nmundo  ",
    )

    processed = preprocess_image_for_ocr(image)
    text, warnings = run_ocr(processed)

    assert processed.getpixel((0, 0)) == 0
    assert processed.getpixel((1, 0)) == 255
    assert text == "hola\nmundo"
    assert warnings == []


def test_run_ocr_handles_missing_tesseract(monkeypatch: pytest.MonkeyPatch) -> None:
    image = Image.new("RGB", (1, 1), color="white")

    monkeypatch.setattr(
        "business_bridge.adapters.ocr_extractors.pytesseract.image_to_string",
        lambda img, config=None: (_ for _ in ()).throw(TesseractNotFoundError()),
    )

    text, warnings = run_ocr(image)

    assert text == ""
    assert warnings == ["Tesseract executable not found; OCR was skipped."]


def test_run_ocr_handles_generic_exceptions(monkeypatch: pytest.MonkeyPatch) -> None:
    image = Image.new("RGB", (1, 1), color="white")

    monkeypatch.setattr(
        "business_bridge.adapters.ocr_extractors.pytesseract.image_to_string",
        lambda img, config=None: (_ for _ in ()).throw(RuntimeError("boom")),
    )

    text, warnings = run_ocr(image)

    assert text == ""
    assert warnings == ["OCR failed: boom"]


def test_extract_text_from_pdf_editable_branch(tmp_path: Path) -> None:
    pdf_path = tmp_path / "editable.pdf"
    document = oe.fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Linea 1\nLinea 2")
    document.save(str(pdf_path))
    document.close()

    page_texts, document_type, warnings = extract_text_from_pdf(pdf_path)

    assert document_type == "pdf_editable"
    assert warnings == []
    assert page_texts[0].startswith("Linea 1")


def test_extract_text_from_pdf_scanned_branch(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    pdf_path = tmp_path / "scanned.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 fake")

    monkeypatch.setattr(
        "business_bridge.adapters.ocr_extractors.extract_pdf_page_texts",
        lambda path: ["", ""],
    )
    monkeypatch.setattr(
        "business_bridge.adapters.ocr_extractors.render_pdf_page_to_image",
        lambda page: Image.new("RGB", (1, 1), color="white"),
    )
    monkeypatch.setattr(
        "business_bridge.adapters.ocr_extractors.preprocess_image_for_ocr",
        lambda image: image,
    )
    monkeypatch.setattr(
        "business_bridge.adapters.ocr_extractors.run_ocr",
        lambda image: ("texto ocr", ["warning"]),
    )
    monkeypatch.setattr(
        "business_bridge.adapters.ocr_extractors.fitz.open",
        lambda path: FakePdfDocument([object(), object()]),
    )

    page_texts, document_type, warnings = extract_text_from_pdf(pdf_path)

    assert document_type == "pdf_scanned"
    assert page_texts == ["texto ocr", "texto ocr"]
    assert warnings == ["Page 1: warning", "Page 2: warning", "OCR applied because the PDF did not contain embedded text."]


def test_extract_text_from_image_and_docx(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    image_path = tmp_path / "sample.png"
    Image.new("RGB", (1, 1), color="white").save(image_path)

    monkeypatch.setattr(
        "business_bridge.adapters.ocr_extractors.run_ocr",
        lambda image: ("texto de imagen", ["warning"]),
    )

    image_texts, image_document_type, image_warnings = extract_text_from_image(image_path)

    assert image_texts == ["texto de imagen"]
    assert image_document_type == "image"
    assert image_warnings == ["warning"]

    docx_path = tmp_path / "sample.docx"
    docx = Document()
    docx.add_paragraph(" Primer parrafo ")
    table = docx.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = " Celda 1 "
    table.rows[0].cells[1].text = " Celda 2 "
    docx.save(str(docx_path))

    docx_texts, docx_document_type, docx_warnings = extract_text_from_docx(docx_path)

    assert docx_document_type == "docx"
    assert docx_warnings == []
    assert "Primer parrafo" in docx_texts[0]
    assert "Celda 1 | Celda 2" in docx_texts[0]


def test_extract_text_from_csv_and_xlsx(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    csv_path = tmp_path / "sample.csv"
    csv_path.write_text("col1,col2\nuno,dos\n", encoding="utf-8")
    csv_calls: list[dict[str, object]] = []

    def fake_read_csv(path, **kwargs):
        csv_calls.append(kwargs)
        if kwargs.get("encoding") is None:
            raise UnicodeDecodeError("utf-8", b"", 0, 1, "boom")
        return pd.DataFrame([{"col1": "uno", "col2": "dos"}])

    monkeypatch.setattr("business_bridge.adapters.ocr_extractors.pd.read_csv", fake_read_csv)

    csv_texts, csv_document_type, csv_warnings = extract_text_from_csv(csv_path)

    assert csv_document_type == "csv"
    assert csv_warnings == []
    assert "uno" in csv_texts[0]
    assert len(csv_calls) == 2

    xlsx_path = tmp_path / "sample.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Hoja 1"
    worksheet["A1"] = " Valor 1 "
    worksheet["B1"] = " Valor 2 "
    workbook.save(str(xlsx_path))

    xlsx_texts, xlsx_document_type, xlsx_warnings = extract_text_from_xlsx(xlsx_path)

    assert xlsx_document_type == "xlsx"
    assert xlsx_warnings == []
    assert "[Sheet] Hoja 1" in xlsx_texts[0]
    assert "Valor 1 Valor 2" in xlsx_texts[0]


@pytest.mark.parametrize(
    ("suffix", "patch_name", "patch_value", "expected_document_type"),
    [
        (".pdf", "extract_text_from_pdf", (["pdf"], "pdf_editable", ["warn"]), "pdf_editable"),
        (".png", "extract_text_from_image", (["image"], "image", ["warn"]), "image"),
        (".txt", "read_text_file", "texto plano", "txt"),
        (".csv", "extract_text_from_csv", (["csv"], "csv", ["warn"]), "csv"),
        (".docx", "extract_text_from_docx", (["docx"], "docx", ["warn"]), "docx"),
        (".xlsx", "extract_text_from_xlsx", (["xlsx"], "xlsx", ["warn"]), "xlsx"),
    ],
)
def test_extract_text_from_file_dispatches_by_extension(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    suffix: str,
    patch_name: str,
    patch_value,
    expected_document_type: str,
) -> None:
    file_path = tmp_path / f"sample{suffix}"
    file_path.write_text("dummy", encoding="utf-8")

    if patch_name == "read_text_file":
        monkeypatch.setattr(
            "business_bridge.adapters.ocr_extractors.read_text_file",
            lambda path: patch_value,
        )
    else:
        monkeypatch.setattr(
            f"business_bridge.adapters.ocr_extractors.{patch_name}",
            lambda path: patch_value,
        )

    page_texts, file_type, document_type, warnings = extract_text_from_file(file_path)

    assert file_type == suffix.lstrip(".")
    assert document_type == expected_document_type
    assert page_texts[0]
    if patch_name == "read_text_file":
        assert warnings == []
    else:
        assert warnings == ["warn"]


def test_extract_text_from_file_handles_unsupported_and_errors(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    xlsb_path = tmp_path / "sample.xlsb"
    xlsb_path.write_text("dummy", encoding="utf-8")
    unsupported_path = tmp_path / "sample.bin"
    unsupported_path.write_text("dummy", encoding="utf-8")
    broken_txt_path = tmp_path / "broken.txt"
    broken_txt_path.write_text("dummy", encoding="utf-8")

    assert extract_text_from_file(xlsb_path)[2] == "unsupported_xlsb"
    assert extract_text_from_file(unsupported_path)[2] == "unsupported"

    monkeypatch.setattr(
        "business_bridge.adapters.ocr_extractors.read_text_file",
        lambda path: (_ for _ in ()).throw(RuntimeError("boom")),
    )

    page_texts, file_type, document_type, warnings = extract_text_from_file(broken_txt_path)

    assert page_texts == [""]
    assert file_type == "txt"
    assert document_type == "error_txt"
    assert warnings[0].startswith("Extraction failed for .txt files:")
