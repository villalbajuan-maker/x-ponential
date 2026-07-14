from __future__ import annotations

from pathlib import Path
from typing import List, Tuple

import fitz
import pandas as pd
import pytesseract
from docx import Document
from openpyxl import load_workbook
from PIL import Image
from pytesseract import TesseractNotFoundError

from business_bridge.adapters.ocr_text import clean_value, normalize_text, read_text_file


SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
UNSUPPORTED_EXTENSIONS = {".xlsb"}


def extract_pdf_page_texts(path: Path) -> List[str]:
    """Extraer el texto incrustado de cada pagina de un PDF."""

    page_texts: List[str] = []
    with fitz.open(path) as pdf_document:
        for page in pdf_document:
            page_texts.append(page.get_text("text") or "")
    return page_texts


def render_pdf_page_to_image(page: fitz.Page) -> Image.Image:
    """Renderizar una pagina PDF como imagen raster para el fallback de OCR."""

    pixmap = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
    return Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)


def preprocess_image_for_ocr(image: Image.Image) -> Image.Image:
    """Convertir una imagen a una version blanco y negro de alto contraste."""

    grayscale = image.convert("L")
    thresholded = grayscale.point(lambda pixel: 255 if pixel > 180 else 0)
    return thresholded


def run_ocr(image: Image.Image) -> Tuple[str, List[str]]:
    """Ejecutar Tesseract OCR y recolectar advertencias recuperables."""

    warnings: List[str] = []
    try:
        text = pytesseract.image_to_string(image, config="--psm 6")
    except TesseractNotFoundError:
        warnings.append("Tesseract executable not found; OCR was skipped.")
        return "", warnings
    except Exception as exc:  # pragma: no cover - safety net for OCR engines
        warnings.append(f"OCR failed: {exc}")
        return "", warnings

    return normalize_text(text), warnings


def extract_text_from_pdf(path: Path) -> Tuple[List[str], str, List[str]]:
    """Extraer texto directo del PDF o rasterizar y aplicar OCR cuando haga falta."""

    warnings: List[str] = []
    page_texts = extract_pdf_page_texts(path)
    combined_text = "\n".join(page_texts).strip()

    if combined_text:
        return [normalize_text(text) for text in page_texts], "pdf_editable", warnings

    ocr_page_texts: List[str] = []
    with fitz.open(path) as pdf_document:
        for page_number, page in enumerate(pdf_document, start=1):
            image = render_pdf_page_to_image(page)
            processed_image = preprocess_image_for_ocr(image)
            page_text, ocr_warnings = run_ocr(processed_image)
            warnings.extend([f"Page {page_number}: {warning}" for warning in ocr_warnings])
            if page_text:
                ocr_page_texts.append(page_text)
            else:
                ocr_page_texts.append("")

    warnings.append("OCR applied because the PDF did not contain embedded text.")
    return ocr_page_texts, "pdf_scanned", warnings


def extract_text_from_image(path: Path) -> Tuple[List[str], str, List[str]]:
    """Aplicar OCR a una sola imagen despues de un pequeno preprocesado."""

    with Image.open(path) as image:
        processed_image = preprocess_image_for_ocr(image)
        text, warnings = run_ocr(processed_image)
    return [text], "image", warnings


def extract_text_from_docx(path: Path) -> Tuple[List[str], str, List[str]]:
    """Tomar texto de parrafos y tablas DOCX y unirlo en un flujo normalizado."""

    warnings: List[str] = []
    document = Document(str(path))
    lines: List[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_values:
                lines.append(" | ".join(row_values))

    return [normalize_text("\n".join(lines))], "docx", warnings


def extract_text_from_csv(path: Path) -> Tuple[List[str], str, List[str]]:
    """Leer un CSV y aplanar el dataframe en texto normalizado."""

    warnings: List[str] = []
    try:
        dataframe = pd.read_csv(path, dtype=str, keep_default_na=False, sep=None, engine="python")
    except Exception:
        dataframe = pd.read_csv(
            path, dtype=str, keep_default_na=False, encoding="latin-1", sep=None, engine="python"
        )

    text = dataframe.to_csv(index=False)
    return [normalize_text(text)], "csv", warnings


def extract_text_from_xlsx(path: Path) -> Tuple[List[str], str, List[str]]:
    """Leer cada hoja de un XLSX como lineas de texto plano."""

    warnings: List[str] = []
    workbook = load_workbook(path, data_only=True, read_only=True)
    lines: List[str] = []

    for worksheet in workbook.worksheets:
        lines.append(f"[Sheet] {worksheet.title}")
        for row in worksheet.iter_rows(values_only=True):
            row_values = [
                clean_value(cell) for cell in row if cell is not None and clean_value(cell) != ""
            ]
            if row_values:
                lines.append("\t".join(row_values))

    workbook.close()
    return [normalize_text("\n".join(lines))], "xlsx", warnings


def extract_text_from_file(path: Path) -> Tuple[List[str], str, str, List[str]]:
    """Despachar la extraccion por extension y devolver paginas, tipo, modo y advertencias."""

    ext = path.suffix.lower()
    warnings: List[str] = []
    file_type = ext.lstrip(".") or "unknown"

    try:
        if ext == ".pdf":
            page_texts, document_type, pdf_warnings = extract_text_from_pdf(path)
            warnings.extend(pdf_warnings)
            return page_texts, file_type, document_type, warnings

        if ext in SUPPORTED_IMAGE_EXTENSIONS:
            page_texts, document_type, image_warnings = extract_text_from_image(path)
            warnings.extend(image_warnings)
            return page_texts, file_type, document_type, warnings

        if ext == ".txt":
            return [normalize_text(read_text_file(path))], file_type, "txt", warnings

        if ext == ".csv":
            page_texts, document_type, csv_warnings = extract_text_from_csv(path)
            warnings.extend(csv_warnings)
            return page_texts, file_type, document_type, warnings

        if ext == ".docx":
            page_texts, document_type, docx_warnings = extract_text_from_docx(path)
            warnings.extend(docx_warnings)
            return page_texts, file_type, document_type, warnings

        if ext == ".xlsx":
            page_texts, document_type, xlsx_warnings = extract_text_from_xlsx(path)
            warnings.extend(xlsx_warnings)
            return page_texts, file_type, document_type, warnings

        if ext in UNSUPPORTED_EXTENSIONS:
            warnings.append("XLSB files are not supported in this flow yet.")
            return [""], file_type, "unsupported_xlsb", warnings

        warnings.append(f"Unsupported file type: {ext.lstrip('.') or 'unknown'}.")
        return [""], file_type, "unsupported", warnings
    except Exception as exc:
        warnings.append(f"Extraction failed for .{file_type} files: {exc}")
        return [""], file_type, f"error_{file_type}", warnings
