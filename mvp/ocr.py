from __future__ import annotations

import re
import unicodedata
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple
from uuid import uuid4

import fitz
import pandas as pd
import pytesseract
from docx import Document
from openpyxl import load_workbook
from PIL import Image
from pytesseract import TesseractNotFoundError


# Flujo OCR del MVP: deteccion de archivos, extraccion de texto, normalizacion
# y deteccion heuristica de campos para documentos de licitacion.
SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
UNSUPPORTED_EXTENSIONS = {".xlsb"}

TENDER_EXPECTED_FIELDS = [
    "tender.process_number",
    "tender.entity",
    "tender.offer_value",
    "tender.execution_term",
    "tender.date",
]

REUSABLE_FIELD_KEYS = {"company.nit", "company.email", "company.phone"}


def detect_file_type(filename: str) -> str:
    """Inferir una etiqueta simple de tipo de archivo desde la extension."""
    suffix = Path(filename).suffix.lower().lstrip(".")
    return suffix or "unknown"


def normalize_text(text: str) -> str:
    """Normalizar Unicode y espacios para que el parseo vea una forma estable."""
    if not text:
        return ""
    normalized = unicodedata.normalize("NFKC", text)
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n").replace("\xa0", " ")
    normalized = re.sub(r"-\n(?=\w)", "", normalized)
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    normalized = "\n".join(line.strip() for line in normalized.split("\n"))
    return normalized.strip()


def safe_preview(text: str, limit: int = 1000) -> str:
    """Devolver una vista previa compacta y limpia de un texto largo."""
    cleaned = normalize_text(text)
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[:limit].rstrip() + "..."


def clean_value(value: Any) -> str:
    """Reducir un valor escalar a una representacion segura de una sola linea."""
    return re.sub(r"\s+", " ", str(value)).strip(" \t\n\r;,.:-")


def read_text_file(path: Path) -> str:
    """Leer un archivo de texto probando UTF-8 primero y Latin-1 como respaldo."""
    try:
        return path.read_text(encoding="utf-8-sig")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1", errors="ignore")


def extract_pdf_page_texts(path: Path) -> List[str]:
    """Extraer el texto incrustado de cada pagina de un PDF."""
    page_texts: List[str] = []
    with fitz.open(path) as pdf_document:
        for page in pdf_document:
            page_texts.append(page.get_text("text") or "")
    return page_texts


def render_pdf_page_to_image(page: fitz.Page) -> Image.Image:
    """Renderizar una pagina PDF como imagen raster para el fallback de OCR."""
    pixmap = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
    return Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)


def preprocess_image_for_ocr(image: Image.Image) -> Image.Image:
    """Convertir una imagen a una version blanco y negro de alto contraste."""
    grayscale = image.convert("L")
    thresholded = grayscale.point(lambda pixel: 255 if pixel > 180 else 0)
    return thresholded


def run_ocr(image: Image.Image) -> Tuple[str, List[str]]:
    """Ejecutar Tesseract OCR y recolectar advertencias recuperables."""
    warnings: List[str] = []
    try:
        text = pytesseract.image_to_string(image, config="--psm 6")
    except TesseractNotFoundError:
        warnings.append("Tesseract executable not found; OCR was skipped.")
        return "", warnings
    except Exception as exc:  # pragma: no cover - safety net for OCR engines
        warnings.append(f"OCR failed: {exc}")
        return "", warnings

    return normalize_text(text), warnings


def extract_text_from_pdf(path: Path) -> Tuple[List[str], str, List[str]]:
    """Extraer texto directo del PDF o rasterizar y aplicar OCR cuando haga falta."""
    warnings: List[str] = []
    page_texts = extract_pdf_page_texts(path)
    combined_text = "\n".join(page_texts).strip()

    if combined_text:
        return [normalize_text(text) for text in page_texts], "pdf_editable", warnings

    # Solo hacemos OCR cuando el PDF no expone texto incrustado.
    ocr_page_texts: List[str] = []
    with fitz.open(path) as pdf_document:
        for page_number, page in enumerate(pdf_document, start=1):
            # Procesamos pagina por pagina para conservar el orden y reportar advertencias precisas.
            image = render_pdf_page_to_image(page)
            processed_image = preprocess_image_for_ocr(image)
            page_text, ocr_warnings = run_ocr(processed_image)
            warnings.extend([f"Page {page_number}: {warning}" for warning in ocr_warnings])
            if page_text:
                ocr_page_texts.append(page_text)
            else:
                ocr_page_texts.append("")

    # Avisamos que el PDF paso por OCR porque no tenia texto embebido.
    warnings.append("OCR applied because the PDF did not contain embedded text.")
    return ocr_page_texts, "pdf_scanned", warnings


def extract_text_from_image(path: Path) -> Tuple[List[str], str, List[str]]:
    """Aplicar OCR a una sola imagen despues de un pequeno preprocesado."""
    image = Image.open(path)
    processed_image = preprocess_image_for_ocr(image)
    text, warnings = run_ocr(processed_image)
    return [text], "image", warnings


def extract_text_from_docx(path: Path) -> Tuple[List[str], str, List[str]]:
    """Tomar texto de parrafos y tablas DOCX y unirlo en un flujo normalizado."""
    warnings: List[str] = []
    document = Document(path)
    lines: List[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_values:
                lines.append(" | ".join(row_values))

    return [normalize_text("\n".join(lines))], "docx", warnings


def extract_text_from_csv(path: Path) -> Tuple[List[str], str, List[str]]:
    """Leer un CSV y aplanar el dataframe en texto normalizado."""
    warnings: List[str] = []
    try:
        dataframe = pd.read_csv(path, dtype=str, keep_default_na=False, sep=None, engine="python")
    except Exception:
        dataframe = pd.read_csv(path, dtype=str, keep_default_na=False, encoding="latin-1", sep=None, engine="python")

    text = dataframe.to_csv(index=False)
    return [normalize_text(text)], "csv", warnings


def extract_text_from_xlsx(path: Path) -> Tuple[List[str], str, List[str]]:
    """Leer cada hoja de un XLSX como lineas de texto plano."""
    warnings: List[str] = []
    workbook = load_workbook(path, data_only=True, read_only=True)
    lines: List[str] = []

    for worksheet in workbook.worksheets:
        lines.append(f"[Sheet] {worksheet.title}")
        for row in worksheet.iter_rows(values_only=True):
            # Saltamos celdas vacias para que el flujo de texto siga compacto y legible.
            row_values = [clean_value(cell) for cell in row if cell is not None and clean_value(cell) != ""]
            if row_values:
                lines.append("\t".join(row_values))

    workbook.close()
    return [normalize_text("\n".join(lines))], "xlsx", warnings


def extract_text_from_file(path: Path) -> Tuple[List[str], str, str, List[str]]:
    """Despachar la extraccion por extension y devolver paginas, tipo, modo y advertencias."""
    ext = path.suffix.lower()
    warnings: List[str] = []
    file_type = ext.lstrip(".") or "unknown"

    try:
        if ext == ".pdf":
            page_texts, document_type, pdf_warnings = extract_text_from_pdf(path)
            warnings.extend(pdf_warnings)
            return page_texts, file_type, document_type, warnings

        if ext in SUPPORTED_IMAGE_EXTENSIONS:
            page_texts, document_type, image_warnings = extract_text_from_image(path)
            warnings.extend(image_warnings)
            return page_texts, file_type, document_type, warnings

        if ext == ".txt":
            return [normalize_text(read_text_file(path))], file_type, "txt", warnings

        if ext == ".csv":
            page_texts, document_type, csv_warnings = extract_text_from_csv(path)
            warnings.extend(csv_warnings)
            return page_texts, file_type, document_type, warnings

        if ext == ".docx":
            page_texts, document_type, docx_warnings = extract_text_from_docx(path)
            warnings.extend(docx_warnings)
            return page_texts, file_type, document_type, warnings

        if ext == ".xlsx":
            page_texts, document_type, xlsx_warnings = extract_text_from_xlsx(path)
            warnings.extend(xlsx_warnings)
            return page_texts, file_type, document_type, warnings

        if ext in UNSUPPORTED_EXTENSIONS:
            # XLSB se deja para despues porque este MVP aun no lo parsea.
            warnings.append("XLSB files are not supported in this MVP yet.")
            return [""], file_type, "unsupported_xlsb", warnings

        warnings.append(f"Unsupported file type: {ext.lstrip('.') or 'unknown'}.")
        return [""], file_type, "unsupported", warnings
    except Exception as exc:
        warnings.append(f"Extraction failed for .{file_type} files: {exc}")
        return [""], file_type, f"error_{file_type}", warnings


def clean_extracted_pages(page_texts: Sequence[str]) -> List[str]:
    """Normalizar cada pagina extraida antes de que la inspeccionen las reglas regex."""
    return [normalize_text(text) for text in page_texts]


def add_detected_item(
    items: List[Dict[str, Any]],
    seen_values: set[Tuple[str, str]],
    field_key: str,
    label: str,
    value: str,
    confidence: float,
    page: Optional[int],
    reusable: bool = False,
    source: str = "regex",
) -> None:
    """Agregar un campo detectado solo si pasa validaciones basicas."""
    cleaned_value = clean_value(value)
    if not cleaned_value:
        return
    # Este filtro rapido elimina falsos positivos obvios antes de llegar a revision.
    digit_count = len(re.sub(r"\D", "", cleaned_value))
    if field_key == "company.phone" and digit_count < 7:
        return
    if field_key == "company.nit" and digit_count < 5:
        return
    if field_key == "company.phone":
        if digit_count < 10 or digit_count > 15:
            return
        if re.search(r"\b(?:nit|proceso|oferta|valor|precio|fecha)\b", cleaned_value, re.IGNORECASE):
            return
    if field_key == "tender.process_number" and digit_count < 3:
        return
    seen_key = (field_key, cleaned_value.casefold())
    if seen_key in seen_values:
        return
    seen_values.add(seen_key)
    items.append(
        {
            "item_id": f"item_{uuid4().hex}",
            "field_key": field_key,
            "label": label,
            "value": cleaned_value,
            "confidence": round(confidence, 2),
            "source": source,
            "page": page,
            "status": "pending_review",
            "reusable": reusable,
        }
    )


def extract_regex_items(page_texts: Sequence[str]) -> List[Dict[str, Any]]:
    """Encontrar campos estructurados con un conjunto pequeno de reglas regex de alto valor."""
    items: List[Dict[str, Any]] = []
    seen_values: set[Tuple[str, str]] = set()

    # Estas expresiones son la primera heuristica del MVP para campos del documento.
    regex_specs = [
        (
            "company.nit",
            "NIT",
            [
                re.compile(r"\bNIT\s*[:\-]?\s*([0-9][0-9.\-]{4,20})\b", re.IGNORECASE),
                re.compile(r"\b([0-9]{3,15}(?:\.[0-9]{3}){1,3}-[0-9])\b"),
            ],
            0.95,
            True,
        ),
        (
            "company.email",
            "Correo",
            [re.compile(r"\b([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})\b")],
            0.98,
            True,
        ),
        (
            "company.phone",
            "Telefono",
            [
                re.compile(r"(?:tel[e\u00e9]fono|celular|mov[i\u00ed]l|phone)\s*[:\-]?\s*([+()0-9\s.-]{7,20})", re.IGNORECASE),
            ],
            0.9,
            True,
        ),
        (
            "tender.offer_value",
            "Valor posible",
            [
                re.compile(
                    r"(?:valor(?: total)?(?: de la oferta)?|oferta|precio)\s*[:\-]?\s*((?:COP\s*)?\$?\s*\d{1,3}(?:[.\s]\d{3})+(?:,\d{1,2})?)",
                    re.IGNORECASE,
                ),
            ],
            0.9,
            False,
        ),
        (
            "tender.date",
            "Fecha",
            [
                re.compile(r"\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b"),
                re.compile(r"\b(\d{4}[/-]\d{1,2}[/-]\d{1,2})\b"),
                re.compile(
                    r"\b(\d{1,2}\s+de\s+(?:enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|setiembre|octubre|noviembre|diciembre)\s+de\s+\d{4})\b",
                    re.IGNORECASE,
                ),
            ],
            0.9,
            False,
        ),
        (
            "tender.process_number",
            "Numero de proceso",
            [
                re.compile(
                    r"(?:proceso(?: de selecci[o\u00f3]n)?|n[u\u00fa]mero de proceso|radicado|no\.?|nro\.?)\s*[:\-]?\s*([A-Z0-9][A-Z0-9\-\/._]{2,80})",
                    re.IGNORECASE,
                ),
                re.compile(r"\b(?:SECOP|LP|MC|CD|CM|SA|IP)\s*[-/ ]?[A-Z0-9][A-Z0-9\-\/._]{2,40}\b", re.IGNORECASE),
            ],
            0.86,
            False,
        ),
        (
            "tender.entity",
            "Entidad",
            [
                re.compile(
                    r"(?:entidad(?: contratante)?|nombre de la entidad|entidad responsable)\s*[:\-]\s*([^\n\r]{3,120})",
                    re.IGNORECASE,
                )
            ],
            0.82,
            False,
        ),
        (
            "tender.execution_term",
            "Plazo de ejecucion",
            [
                re.compile(
                    r"(?:plazo de ejecuci[o\u00f3]n|duraci[o\u00f3]n|tiempo de ejecuci[o\u00f3]n)\s*[:\-]?\s*([^\n\r]{3,80})",
                    re.IGNORECASE,
                ),
                re.compile(
                    r"\bplazo\b\s*[:\-]?\s*([0-9]{1,3}\s*(?:d[i\u00ed]as?|meses?|semanas?)[^\n\r]*)",
                    re.IGNORECASE,
                ),
            ],
            0.8,
            False,
        ),
    ]

    # Procesamos cada pagina por separado para poder deduplicar sin perder contexto de pagina.
    for page_index, page_text in enumerate(page_texts, start=1):
        text = normalize_text(page_text)
        if not text:
            continue
        for field_key, label, patterns, confidence, reusable in regex_specs:
            for pattern in patterns:
                for match in pattern.finditer(text):
                    candidate = match.group(1) if match.groups() else match.group(0)
                    add_detected_item(
                        items=items,
                        seen_values=seen_values,
                        field_key=field_key,
                        label=label,
                        value=candidate,
                        confidence=confidence,
                        page=page_index if len(page_texts) > 1 else None,
                        reusable=reusable or field_key in REUSABLE_FIELD_KEYS,
                    )

    return items


def calculate_overall_confidence(detected_items: Sequence[Dict[str, Any]], document_type: str, raw_text: str) -> float:
    """Colapsar la confianza por item en una sola puntuacion de documento."""
    if detected_items:
        average_confidence = sum(float(item.get("confidence", 0.0)) for item in detected_items) / len(detected_items)
        if document_type in {"pdf_scanned", "image"}:
            average_confidence -= 0.05
        return round(max(0.0, min(0.99, average_confidence)), 2)
    if raw_text.strip():
        return 0.45 if document_type in {"pdf_scanned", "image"} else 0.5
    return 0.0


def build_extraction_record(stored_file: Path, document_id: Optional[str] = None) -> Dict[str, Any]:
    """Armar el registro persistido que usa el flujo de revision del frontend."""
    stored_file = Path(stored_file)
    if document_id is None:
        document_id = stored_file.name.split("__", 1)[0]
    original_file_name = stored_file.name.split("__", 1)[1] if "__" in stored_file.name else stored_file.name
    page_texts, file_type, document_type, warnings = extract_text_from_file(stored_file)
    cleaned_page_texts = clean_extracted_pages(page_texts)
    combined_text = "\n".join(text for text in cleaned_page_texts if text).strip()
    detected_items = extract_regex_items(cleaned_page_texts)
    # Lo que no detecte la regex se vuelve un faltante explicito para la UI.
    detected_field_keys = {item["field_key"] for item in detected_items}
    missing_fields = [field_key for field_key in TENDER_EXPECTED_FIELDS if field_key not in detected_field_keys]

    if not combined_text and not warnings:
        warnings.append("No text could be extracted from the document.")
    elif not combined_text:
        warnings.append("Extraction completed but no usable text was found.")

    return {
        "document_id": document_id,
        "file_name": original_file_name,
        "file_type": file_type,
        "document_type": document_type,
        "confidence": calculate_overall_confidence(detected_items, document_type, combined_text),
        "raw_text_preview": safe_preview(combined_text),
        "detected_items": detected_items,
        "missing_fields": missing_fields,
        "warnings": warnings,
        "supplemental_answers": {},
    }
